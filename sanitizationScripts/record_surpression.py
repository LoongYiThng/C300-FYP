import sys
from docx import Document
import re

if sys.argv[1]=="document":
    regex={
        "email": "\S+@(\S+|\.\S+)",
        "age": "\d+\syears\sold",
        "money": "\$\d+(?:\,\d+|\d+)+(?:\.\d+)?",
        "phone number": "(?:\+65|\+65\s)?\d{4}\s?\d{4}",
        "name or place": "[A-Z][a-z]+\s?"
    }
    sensitive_types=[]

    # define sensitive data
    def is_sensitive(sentence):
        sensitive=False
        for key in regex:
            if re.search(regex[key], sentence):
                sensitive=True
                sensitive_types.append(key)
        return sensitive

    def sanitize(sentence, sensitive_types:list):
        for i in sensitive_types:
            sentence=re.sub(regex[i], "", sentence)
        return sentence

    # open document
    doc = Document("test.docx")

    # sanitize
    sanitized_doc=Document()
    output_paragraph=""

    for paragraph in doc.paragraphs:
        original_text=paragraph.text
        sentences=re.split("\.\s", original_text)
        for sentence in sentences:
            if is_sensitive(sentence):
                output_paragraph+=sanitize(sentence, sensitive_types)+". "
            else:
                output_paragraph+=sentence
            sensitive_types=[]

        sanitized_doc.add_paragraph(output_paragraph)
        output_paragraph=""

    # output
    sanitized_doc.save("surpressed.docx")
    print("Done! check save folder")

elif sys.argv[1]=="excel":
    pass