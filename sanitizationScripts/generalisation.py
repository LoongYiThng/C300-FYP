import sys
from docx import Document
import re

if sys.argv[1]=="document":
    regex={
        "age": "\d+\syears\sold",
        "money": "\$\d+(?:\,\d+|\d+)+(?:\.\d+)?",
        "name or place": "[A-Z][a-z]+\s?"
    }
    sensitive_types=[]

    # define sensitive data
    def is_sensitive(sentence):
        sensitive=False
        for key in regex:
            if re.search(regex[key], sentence):
                sensitive=True
                sensitive_types.append(key)
        return sensitive

    def sanitize(sentence, sensitive_types:list):
        if "age" in sensitive_types:
            all_instances=re.findall(regex["age"], sentence)
            for matched in all_instances:
                age=re.search("\d+", matched)
                age=int(age.group())
                sentence=re.sub(re.escape(matched), str(age-10)+" - "+str(age+10), sentence)

        if "money" in sensitive_types:
            all_instances=re.findall(regex["money"], sentence)
            for matched in all_instances:
                money=float(re.sub("\$|,", "", matched))
                sentence=re.sub(re.escape(matched), str(money-money*0.2)+" - "+str(money+money*0.2), sentence)

        if "name or place" in sensitive_types:
            pass

        return sentence

    # open document
    doc = Document("sanitizationScripts/test.docx")

    # sanitize
    sanitized_doc=Document()
    output_paragraph=""

    for paragraph in doc.paragraphs:
        original_text=paragraph.text
        sentences=re.split("\.\s", original_text)
        for sentence in sentences:
            if is_sensitive(sentence):
                output_paragraph+=sanitize(sentence, sensitive_types)+". "
            else:
                output_paragraph+=sentence
            sensitive_types=[]

        sanitized_doc.add_paragraph(output_paragraph)
        output_paragraph=""

    # output
    sanitized_doc.save("sanitizationScripts/generalized.docx")
    print("Done! check save folder")

elif sys.argv[1]=="excel":
    pass